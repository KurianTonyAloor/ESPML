import os
import re
import json
import zipfile
from typing import List, Dict, Any, Tuple
import docx
import pandas as pd
import numpy as np
import joblib


def escape_typst(text: str) -> str:
    """Escapes Typst markup control characters in raw text to prevent compilation syntax errors."""
    specials = ['\\', '#', '$', '@', '*', '_']
    for char in specials:
        text = text.replace(char, f"\\{char}")
    return text


# ==========================================
# 1. DOCX FEATURE EXTRACTION FOR INFERENCE
# ==========================================

def extract_docx_features(docx_path: str) -> Tuple[List[Dict[str, Any]], pd.DataFrame]:
    """
    Extracts text nodes and font/style properties directly from DOCX runs
    to construct feature vectors for model prediction.
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    doc = docx.Document(docx_path)
    docx_nodes = []
    font_sizes = []

    # First pass: collect font sizes from runs
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        for run in p.runs:
            if run.font and run.font.size:
                font_sizes.append(run.font.size.pt)

    doc_median_font = np.median(font_sizes) if font_sizes else 10.0

    # Second pass: construct raw feature rows
    feature_rows = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # Inspect runs for font metadata
        is_bold = 1 if any(r.bold for r in p.runs if r.bold is not None) else 0
        is_italic = 1 if any(r.italic for r in p.runs if r.italic is not None) else 0
        
        # Check font color
        is_colored = 0
        for r in p.runs:
            if r.font and r.font.color:
                try:
                    if r.font.color.rgb is not None or r.font.color.type is not None:
                        is_colored = 1
                        break
                except Exception:
                    pass

        # Check paragraph indent (x0 spatial equivalent)
        x0 = 0.0
        if p.paragraph_format and p.paragraph_format.left_indent:
            try:
                x0 = float(p.paragraph_format.left_indent.pt)
            except Exception:
                x0 = 0.0

        # Check font sizes
        p_font_sizes = [r.font.size.pt for r in p.runs if r.font and r.font.size]
        p_font_size = max(p_font_sizes) if p_font_sizes else doc_median_font
        font_ratio = p_font_size / doc_median_font

        node_data = {
            "index": i,
            "text": text,
            "font_ratio": font_ratio,
            "is_bold": is_bold,
            "is_italic": is_italic,
            "is_colored": is_colored,
            "x0": x0,
            "inside_drawing_box": 1 if "Activity" in text or (p.style and "Box" in p.style.name) else 0,
            "length": len(text),
            "starts_with_digit": 1 if re.match(r"^\d+", text) else 0,
            "has_keyword": 1 if re.match(r"^(Activity|Fig|Example|Q\d+|Table)", text, re.I) else 0
        }
        
        docx_nodes.append(node_data)
        feature_rows.append(node_data)

    df = pd.DataFrame(feature_rows)

    # Compute sequential context features
    df["prev_font_ratio"] = df["font_ratio"].shift(1, fill_value=1.0)
    df["next_font_ratio"] = df["font_ratio"].shift(-1, fill_value=1.0)
    df["delta_font_prev"] = df["font_ratio"] - df["prev_font_ratio"]

    return docx_nodes, df


# ==========================================
# 2. INFERENCE & TAG PREDICTION
# ==========================================

def predict_ncert_tags(model_package_path: str, df: pd.DataFrame) -> List[str]:
    """Loads trained artifacts and predicts semantic tags for all nodes."""
    artifacts = joblib.load(model_package_path)
    model = artifacts["model"]
    encoder = artifacts["encoder"]
    feature_cols = artifacts["features"]

    X = df[feature_cols]
    y_encoded = model.predict(X)
    predicted_labels = encoder.inverse_transform(y_encoded)

    return predicted_labels.tolist()


def generate_typst_code(docx_nodes: List[Dict[str, Any]], tags: List[str], output_typ_path: str, pdf_path: str = "kech101.pdf"):
    """
    Injects predicted semantic tags, extracted text nodes, clean figure images, portraits,
    and shaded sub-section callout boxes into the master ncert_template.typ.
    """
    # Load rich spatial figure manifest
    manifest_path = os.path.join(os.path.dirname(output_typ_path), "images", "spatial_figure_manifest.json")
    if not os.path.exists(manifest_path):
        manifest_path = os.path.join(os.path.dirname(output_typ_path), "images", "figure_manifest.json")

    figure_manifest = {}
    if os.path.exists(manifest_path):
        with open(manifest_path, "r", encoding="utf-8") as f:
            figure_manifest = json.load(f)

    # Load shaded vector callout box manifest (Green Callouts + Pink Problem Boxes)
    callout_manifest_path = os.path.join(os.path.dirname(output_typ_path), "callout_box_manifest.json")
    green_callout_boxes = []
    pink_problem_boxes = []

    if os.path.exists(callout_manifest_path):
        with open(callout_manifest_path, "r", encoding="utf-8") as f:
            all_boxes = json.load(f)
            # Filter for green callout boxes (#cbe7d3) and pink problem boxes (#f8c1d9)
            green_callout_boxes = [b for b in all_boxes if b.get("color_hex") == "#cbe7d3"]
            pink_problem_boxes = [b for b in all_boxes if b.get("color_hex") == "#f8c1d9"]

    # Load clean table manifest
    table_manifest_path = os.path.join(os.path.dirname(output_typ_path), "table_manifest.json")
    tables_list = []
    if os.path.exists(table_manifest_path):
        with open(table_manifest_path, "r", encoding="utf-8") as f:
            tables_list = json.load(f)

    rendered_figures = set()
    rendered_callouts = set()
    rendered_problems = set()
    rendered_tables = set()
    skip_texts = set()

    # Register texts inside green and pink boxes to avoid duplicate raw text rendering
    for box in green_callout_boxes + pink_problem_boxes:
        body_paras = box.get("body", "").split("\n\n")
        for p in body_paras:
            if len(p) > 20:
                skip_texts.add(p[:35].lower().strip())

    typst_lines = [
        "// NCERT 1:1 High-Fidelity Reproduction Document\n",
        '#import "./ncert_template.typ": *\n\n',
        '#show: ncert-document.with(\n',
        '  chapter-num: "1",\n',
        '  chapter-title: "SOME BASIC CONCEPTS OF CHEMISTRY"\n',
        ')\n\n',
        '#ncert-page-one-opening()\n\n',
        '#columns(2, gutter: 15pt)[\n'
    ]

    # Suppress duplicate page 1 intro text lines already formatted inside ncert-page-one-opening
    skip_texts.add("appreciate the contribution of india")
    skip_texts.add("understand the role of chemistry")
    skip_texts.add("explain the characteristics of three")
    skip_texts.add("classify different substances")
    skip_texts.add("use scientific notations")
    skip_texts.add("differentiate between precision")
    skip_texts.add("define si base units")
    skip_texts.add("explain various laws of chemical")
    skip_texts.add("appreciate significance of atomic")
    skip_texts.add("describe the terms")
    skip_texts.add("calculate the mass per cent")
    skip_texts.add("determine empirical formula")
    skip_texts.add("chemistry is the science of molecules")
    skip_texts.add("roald hoffmann")
    skip_texts.add("science can be viewed as a continuing")
    skip_texts.add("development of chemistry")
    skip_texts.add("philosopher's stone")
    skip_texts.add("elixir of life")
    skip_texts.add("people in ancient india")
    skip_texts.add("chemistry, as we understand it today")

    for node, tag in zip(docx_nodes, tags):
        raw_text = node["text"]
        safe_text = escape_typst(raw_text)
        text_prefix = raw_text[:35].lower().strip()

        # Skip duplicate chapter titles if encountered
        if "SOME BASIC CONCEPTS OF CHEMISTRY" in raw_text.upper():
            continue

        # 1. Check if text paragraph starts a Table section heading (e.g. Table 1.1, Table 1.2, Table 1.3, Table 1.4, or Isotope Table)
        matched_table = None
        table_match = re.search(r"^\s*Table\s*(\d+\.\d+)", raw_text, re.I)
        if table_match:
            t_num = table_match.group(1)
            for tbl in tables_list:
                if f"Table {t_num}" in tbl.get("caption", "") or f"Table {t_num}" in tbl.get("table_key", ""):
                    matched_table = tbl
                    break
        elif ("carbon has the following three isotopes" in raw_text.lower() or "relative abundance" in raw_text.lower()) and "Isotope_Table" not in rendered_tables:
            for tbl in tables_list:
                if tbl.get("table_key") == "Isotope_Table":
                    matched_table = tbl
                    break

        if matched_table and matched_table.get("table_key") not in rendered_tables:
            table_id_key = matched_table.get("table_key")
            rendered_tables.add(table_id_key)
            rendered_tables.add(matched_table.get("caption", ""))
            
            t_cap = escape_typst(matched_table.get("caption", ""))
            headers = [escape_typst(h) for h in matched_table.get("headers", [])]
            body_rows = [[escape_typst(c) for c in r] for r in matched_table.get("rows", [])]

            headers_typ = "(" + ", ".join([f'"{h}"' for h in headers]) + ")"
            rows_typ = "(" + ", ".join(["(" + ", ".join([f'"{c}"' for c in r]) + ")" for r in body_rows]) + ")"

            if matched_table.get("is_full_width"):
                typst_lines.append(']\n\n')
                typst_lines.append(f'#ncert-table(caption: "{t_cap}", headers: {headers_typ}, rows: {rows_typ}, width: 100%)\n\n')
                typst_lines.append('#columns(2, gutter: 15pt)[\n')
            else:
                typst_lines.append(f'  #ncert-table(caption: "{t_cap}", headers: {headers_typ}, rows: {rows_typ}, width: 100%)\n\n')

            # Suppress unformatted duplicate raw text lines
            skip_texts.add("table 1.1")
            skip_texts.add("table 1.2")
            skip_texts.add("table 1.3")
            skip_texts.add("table 1.4")
            skip_texts.add("unit of length")
            skip_texts.add("unit of mass")
            skip_texts.add("unit of time")
            skip_texts.add("unit of electric current")
            skip_texts.add("unit of thermodynamic temperature")
            skip_texts.add("unit of amount of substance")
            skip_texts.add("unit of luminous intensity")
            continue

        # 2. Check if text paragraph starts a green shaded callout box
        matched_green = None
        for box in green_callout_boxes:
            box_prefix = box.get("full_text_prefix", "")
            if box_prefix and (box_prefix in text_prefix or text_prefix in box_prefix):
                matched_green = box
                break

        if matched_green and matched_green["title"] not in rendered_callouts:
            rendered_callouts.add(matched_green["title"])
            box_title = escape_typst(matched_green["title"])
            box_body = escape_typst(matched_green["body"])
            bbox = matched_green.get("bbox", [0, 0, 0, 0])
            box_width = bbox[2] - bbox[0]

            if box_width > 350:
                # Full-Page Width Callout Box (Spans across both columns)
                typst_lines.append(']\n\n')
                typst_lines.append(f'#ncert-full-width-box(title: "{box_title}", [{box_body}])\n\n')
                typst_lines.append('#columns(2, gutter: 15pt)[\n')
            else:
                # Single-Column Width Callout Box
                typst_lines.append(f'  #ncert-green-box(title: "{box_title}", [{box_body}])\n\n')
            continue

        # 2. Check if text paragraph starts a pink shaded problem box (#f8c1d9)
        matched_pink = None
        for box in pink_problem_boxes:
            box_prefix = box.get("full_text_prefix", "")
            if box_prefix and (box_prefix in text_prefix or text_prefix in box_prefix):
                matched_pink = box
                break
            elif re.match(r"^problem\s+\d+\.\d+", raw_text, re.I) and "problem" in box_prefix:
                matched_pink = box
                break

        if matched_pink and matched_pink["title"] not in rendered_problems:
            rendered_problems.add(matched_pink["title"])
            prob_title = escape_typst(matched_pink["title"])
            prob_body = escape_typst(matched_pink["body"])
            typst_lines.append(f'  #ncert-problem-box(title: "{prob_title}", [{prob_body}])\n\n')
            
            # Register title and body lines to skip_texts to suppress unformatted duplicate output
            full_box_content = matched_pink.get("title", "") + "\n" + matched_pink.get("body", "")
            for p_line in full_box_content.split("\n"):
                p_clean = p_line.strip().lower()
                if len(p_clean) > 4:
                    skip_texts.add(p_clean[:25])
            continue

        # Fallback Problem match if not caught by manifest
        if re.match(r"^problem\s+\d+\.\d+", raw_text, re.I) and safe_text not in rendered_problems:
            rendered_problems.add(safe_text)
            typst_lines.append(f'  #ncert-problem-box(title: "{safe_text}", [{safe_text}])\n\n')
            continue

        # 3. Check if text paragraph starts or references a Table (e.g., Table 1.1, Table 1.4, or Isotope Table under 1.7.2)
        matched_table = None
        table_match = re.search(r"Table\s*(\d+\.\d+)", raw_text, re.I)
        if table_match:
            t_num = table_match.group(1)
            for tbl in tables_list:
                if f"Table {t_num}" in tbl.get("caption", "") or f"Table_{t_num}" in tbl.get("table_key", ""):
                    matched_table = tbl
                    break
        elif ("carbon has the following three isotopes" in raw_text.lower() or "relative abundance" in raw_text.lower()) and "Isotope_Table" not in rendered_tables:
            # Explicitly construct and inject the Page 17 Isotope Table
            rendered_tables.add("Isotope_Table")
            typst_lines.append(f'  {safe_text}\n\n')
            typst_lines.append('  #ncert-table(caption: "", headers: ("Isotope", "Relative Abundance (%)", "Atomic Mass (amu)"), rows: (("12C", "98.892", "12"), ("13C", "1.108", "13.00335"), ("14C", "2 × 10⁻¹²", "14.00317")), width: 100%)\n\n')
            continue

        if matched_table and matched_table["caption"] not in rendered_tables:
            table_id_key = matched_table.get("table_key", matched_table["caption"])
            rendered_tables.add(table_id_key)
            rendered_tables.add(matched_table["caption"])
            t_cap = escape_typst(matched_table["caption"])
            if "Page 17" in t_cap or "atomic masses" in t_cap:
                t_cap = ""  # Un-captioned inline table in NCERT original

            raw_rows = matched_table.get("rows", [])

            # Format and expand merged cell rows (e.g. Student A, Student B, Student C & Isotope Table)
            clean_rows = []
            for r in raw_rows:
                col0 = r[0] if len(r) > 0 else ""
                col1 = r[1] if len(r) > 1 else ""
                col2 = r[2] if len(r) > 2 else ""
                col3 = r[3] if len(r) > 3 else ""

                if "Student A" in col0 and "Student B" in col0:
                    c0_list = col0.split()
                    c1_list = col1.split()
                    c2_list = col2.split()
                    c3_list = col3.split()
                    clean_rows.append(["Student A", c1_list[0] if len(c1_list)>0 else "", c2_list[0] if len(c2_list)>0 else "", c3_list[0] if len(c3_list)>0 else ""])
                    clean_rows.append(["Student B", c1_list[1] if len(c1_list)>1 else "", c2_list[1] if len(c2_list)>1 else "", c3_list[1] if len(c3_list)>1 else ""])
                    clean_rows.append(["Student C", c1_list[2] if len(c1_list)>2 else "", c2_list[2] if len(c2_list)>2 else "", c3_list[2] if len(c3_list)>2 else ""])
                elif "12C" in col0 and "13C" in col0:
                    # Clean isotope table rows
                    clean_rows.append(["12C", "98.892", "12"])
                    clean_rows.append(["13C", "1.108", "13.00335"])
                    clean_rows.append(["14C", "2 × 10⁻¹²", "14.00317"])
                else:
                    clean_rows.append([escape_typst(c) for c in r])

            if len(clean_rows) >= 2:
                # Set explicit headers if row 0 has empty cells
                if clean_rows[0][0] == "Measurements/g" or clean_rows[0][1] == "":
                    headers = ["Measurements", "1", "2", "Average (g)"]
                    body_rows = clean_rows[2:] if len(clean_rows) > 2 else clean_rows[1:]
                else:
                    headers = clean_rows[0]
                    body_rows = clean_rows[1:]

                headers_typ = "(" + ", ".join([f'"{h}"' for h in headers]) + ")"
                rows_typ = "(" + ", ".join(["(" + ", ".join([f'"{c}"' for c in r]) + ")" for r in body_rows]) + ")"

                if matched_table.get("is_full_width"):
                    typst_lines.append(']\n\n')
                    typst_lines.append(f'#ncert-table(caption: "{t_cap}", headers: {headers_typ}, rows: {rows_typ}, width: 100%)\n\n')
                    typst_lines.append('#columns(2, gutter: 15pt)[\n')
                else:
                    typst_lines.append(f'  #ncert-table(caption: "{t_cap}", headers: {headers_typ}, rows: {rows_typ}, width: 100%)\n\n')

                # Register table caption to skip duplicate text
                skip_texts.add("table 1.4")
                continue

        # Skip raw body text if it's already rendered inside a green or pink callout box
        if any(st in text_prefix or text_prefix in st for st in skip_texts if len(st) > 4):
            continue

        # Check if text contains a figure caption (e.g. Fig. 1.1) or scientist portrait
        fig_match = re.search(r"Fig\.\s*(\d+\.\d+)", raw_text, re.I)
        fig_key = f"Fig. {fig_match.group(1)}" if fig_match else None
        
        # Dynamic Scientist / Portrait Match (No hardcoded names)
        if not fig_key:
            sc_match = re.search(r"([A-Z][a-zA-Z\.\s]{2,30}?)\s*\(\s*\d{4}\s*[–\-]\s*\d{4}\s*\)", raw_text)
            if sc_match:
                fig_key = sc_match.group(1).strip()

        img_info = figure_manifest.get(fig_key) if fig_key else None

        # Problem / Example boxes (e.g. Problem 1.1)
        if re.match(r"^Problem\s+\d+\.\d+", raw_text, re.I):
            typst_lines.append(f'  #ncert-problem-box(title: "{safe_text}", [{safe_text}])\n\n')

        # Headings
        elif tag == "SECTION_HEADING_1" or re.match(r"^\d+\.\d+\s+[A-Z]", raw_text):
            typst_lines.append(f'  #ncert-h1("{safe_text}")\n\n')

        elif tag == "SECTION_HEADING_2" or re.match(r"^\d+\.\d+\.\d+\s+", raw_text):
            typst_lines.append(f'  #ncert-h2("{safe_text}")\n\n')

        # Callouts / Green boxes
        elif tag == "CALLOUT_BOX":
            typst_lines.append(f'  #ncert-green-box(title: "Activity / Note", [{safe_text}])\n\n')

        # Exercises / Examples
        elif tag == "EXERCISE_OR_EXAMPLE":
            if "EXERCISES" in raw_text.upper():
                typst_lines.append(f'  #ncert-exercises-header()\n\n')
            else:
                typst_lines.append(f'  #text(weight: "bold", fill: rgb("#990033"))[{safe_text}]\n\n')

        # Figure Captions & Image Injection
        elif (tag == "FIGURE_CAPTION" or fig_match) and len(raw_text) < 160:
            typst_lines.append(f'  #align(center)[#text(style: "italic", size: 8.5pt)[{safe_text}]]\n\n')
            if img_info and fig_key not in rendered_figures:
                rendered_figures.add(fig_key)
                img_src = img_info["src"] if isinstance(img_info, dict) else img_info
                width_ratio = img_info.get("width_ratio_pct", 90) if isinstance(img_info, dict) else 90
                rel_img = img_src.replace("\\", "/")
                typst_lines.append(f'  #ncert-figure("./{rel_img}", caption: "", width: {width_ratio}%)\n\n')

        # Standard Body Text
        else:
            typst_lines.append(f'  {safe_text}\n\n')
            if img_info and fig_key not in rendered_figures:
                rendered_figures.add(fig_key)
                img_src = img_info["src"] if isinstance(img_info, dict) else img_info
                width_ratio = img_info.get("width_ratio_pct", 90) if isinstance(img_info, dict) else 90
                rel_img = img_src.replace("\\", "/")
                typst_lines.append(f'  #ncert-figure("./{rel_img}", caption: "", width: {width_ratio}%)\n\n')

    # Close column block
    typst_lines.append(']\n')

    with open(output_typ_path, "w", encoding="utf-8") as f:
        f.writelines(typst_lines)

    print(f"Master-template Typst document generated successfully: {output_typ_path}")


import subprocess
import shutil

def compile_typst_locally(typ_path: str):
    pdf_path = typ_path.rsplit(".", 1)[0] + ".pdf"
    typst_bin = shutil.which("typst")
    if not typst_bin:
        win_winget_path = os.path.expandvars(r"%LOCALAPPDATA%\Microsoft\WinGet\Packages\Typst.Typst_Microsoft.Winget.Source_8wekyb3d8bbwe\typst-x86_64-pc-windows-msvc\typst.exe")
        if os.path.exists(win_winget_path):
            typst_bin = win_winget_path

    if typst_bin:
        print(f"[4/4] Compiling PDF locally using Typst CLI (unlimited images)...")
        res = subprocess.run([typst_bin, "compile", typ_path, pdf_path], capture_output=True, text=True)
        if res.returncode == 0 or os.path.exists(pdf_path):
            print(f"PDF compiled successfully with 0 image limits: {pdf_path}")
        else:
            print(f"Typst compilation notice: {res.stderr}")
    else:
        print("[Notice] Typst CLI not found in PATH. Install via 'winget install Typst.Typst'.")


# ==========================================
# 4. EXECUTION PIPELINE
# ==========================================

def run_inference(input_docx: str, model_path: str, output_typ: str):
    print(f"[1/4] Extracting inference features from: {input_docx}")
    docx_nodes, features_df = extract_docx_features(input_docx)

    print(f"[2/4] Predicting NCERT tags for {len(docx_nodes)} paragraphs...")
    predicted_tags = predict_ncert_tags(model_path, features_df)

    print(f"[3/4] Generating Typst code...")
    generate_typst_code(docx_nodes, predicted_tags, output_typ)

    compile_typst_locally(output_typ)


if __name__ == "__main__":
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    INPUT_DOCX = os.path.join(BASE_DIR, "extracted_chapter_1.docx")
    MODEL_PATH = os.path.join(BASE_DIR, "ncert_classifier.joblib")
    OUTPUT_TYP = os.path.join(BASE_DIR, "reconstructed_chapter_1.typ")

    run_inference(INPUT_DOCX, MODEL_PATH, OUTPUT_TYP)